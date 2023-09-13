import os
from striprtf.striprtf import rtf_to_text
from docx import Document
import PyPDF2

def rtf_to_txt(rtf_file):
    with open(rtf_file, 'r') as file:
        rtf_text = file.read()
    print(f"{rtf_file} done.")
    return rtf_to_text(rtf_text)

def docx_to_txt(doc_path):
    try:
        doc = Document(doc_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        # Handling tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        print(f"{doc_path} done.")
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error occurred: {e}")
        return ''

def pdf_to_txt(pdf_file):
    pdf_file_obj = open(pdf_file, 'rb')
    pdf_reader = PyPDF2.PdfReader(pdf_file_obj)
    full_text = []
    for page_num in range(len(pdf_reader.pages)):
        page_obj = pdf_reader.pages[page_num]
        full_text.append(page_obj.extract_text())
    print(f"{pdf_file} done.")
    return '\n'.join(full_text)


filePath = '../data/legal_drafts/GST/Registration/GSTR-7.pdf'
pdfText = pdf_to_txt(filePath)
# docxText = docx_to_txt(filePath)
# rtfText = rtf_to_txt(filePath)
with open('../data/legal_drafts_text/GSTR-7.txt', 'w') as outFile:
    outFile.write(pdfText)


# Starting directory
start_dir = "../data/legal_drafts"
# New directory for text files
output_dir = "../data/legal_drafts_text"

# docxFilePath = start_dir + '/Child Custody/Income Affidavit.docx'


rtfCount: int = 0
docxCount: int = 0 
pdfCount: int = 0
for foldername, subfolders, filenames in os.walk(start_dir):
    # print(f"foldername = {foldername}")
    # print(f"subfolders = {subfolders}")
    # print(f"filenames = {filenames}")
    for filename in filenames:
        # print(f"filename = {filename}")
        # Full file path
        file_path = os.path.join(foldername, filename)
        # print(f"file_path = {file_path}")
        # Extension of the file
        _, ext = os.path.splitext(file_path)
        # Depending on the extension, call the appropriate conversion function
        if ext == '.rtf':
            rtfCount += 1
            text = rtf_to_txt(file_path)
            #  # Construct the new file path in the output directory
            # new_folder = foldername.replace(start_dir, output_dir, 1)
            # # Create directories, if they don't exist
            # os.makedirs(new_folder, exist_ok=True)
            # new_file_path = os.path.join(new_folder, filename.split('.')[0] + '.txt')
            # # Write the text to the new file
            # with open(new_file_path, 'w', errors='ignore') as new_file:
            #     new_file.write(text)
        elif ext == '.docx':
            docxCount += 1
            text = docx_to_txt(file_path)
        elif ext == '.pdf':
            pdfCount += 1
            text = pdf_to_txt(file_path)
        else:
            continue
        # Construct the new file path in the output directory
        new_folder = foldername.replace(start_dir, output_dir, 1)
        # Create directories, if they don't exist
        os.makedirs(new_folder, exist_ok=True)
        new_file_path = os.path.join(new_folder, filename.split('.')[0] + '.txt')
        # Write the text to the new file
        with open(new_file_path, 'w', errors='ignore') as new_file:
            new_file.write(text)

print(f"docCount = rtf: {rtfCount} + docx: {docxCount} + pdf: {pdfCount}")