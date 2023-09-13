from docx import Document


# Starting directory
start_dir = "../data/3500_legal_drafts"
# New directory for text files
output_dir = "../data/3500_legal_drafts_text"

doc = Document(start_dir+'/Child Custody/Income Affidavit.docx')
full_text = ''


print(doc.tables)
print(type(doc.tables[0]))
once = 1
for para in doc.paragraphs:
    full_text += para.text + '\n'
for table in doc.tables:
    table_data = ''
    for row in table.rows:
        row_data = ''
        for cell in row.cells:
            row_data += cell.text + '\t\t\t'
        table_data += row_data + '\n'
    
    
        # if once == 1:
        #     with open(output_dir+'/Income Affidavit.txt', 'w') as outFile:
        #         outFile.write(str(table_data))
    full_text += table_data + '\n'

# print("\n\nrow_data {}\n".format(row_data))
# print("\n\ntable_data {}\n".format(table_data))
# print("\n\nfull_text {}\n".format(full_text))
with open(output_dir+'/Income Affidavit-new.txt', 'w') as outFile:
    outFile.write(full_text)